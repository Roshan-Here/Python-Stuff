import docx
from docx import Document
from docx.shared import Pt,Inches # for font sizing...
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor # for font color setting....
from docx.enum.dml import MSO_THEME_COLOR #MSO theme selector color
from docx.enum.text import WD_ALIGN_PARAGRAPH # alignmennt
from docx.enum.style import WD_STYLE

###############################################
# https://github.com/python-openxml/python-docx/issues/105#issuecomment-442786431
#############################################
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom') #top
    bottom.set(qn('w:val'), 'single') #double
    bottom.set(qn('w:sz'), '20') #helps to assign line width 
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

#############################################
#############################################




###############################
# https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
# https://github.com/python-openxml/python-docx/issues/74#issuecomment-244602378
################################

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


#add a hyperlink with a custom color and no underline
# hyperlink = add_hyperlink(p, 'http://www.google.com', 'Google', 'FF8822', False)


###############################
# https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
################################


d = Document()

section = d.sections[0]  # Get the first section of the document
section.left_margin = Inches(.5)  # Set the left margin to 1 inch
section.right_margin = Inches(.5)  # Set the right margin to 1 inch
section.top_margin = Inches(.5)   # Set the top margin to 1 inch
section.bottom_margin = Inches(.5)

p = d.add_heading("adding para",level=0)


para = d.add_paragraph()
# bold para
para.add_run('wow its nice').bold = True
# italic para on same stanzzza
para.add_run(' oopz').italic = True

paragraph = d.add_paragraph('Normal text, ')
run = paragraph.add_run('text with emphasis.' ,'Emphasis')
j = para.add_run(' check ', 'Strong')
# added font name
k = para.add_run('its a Subtle Reference check how i look', 'Subtle Reference')
k.font.name = 'Times New Roman'
# font sizing
k.font.size = Pt(25)
# font underline
k.font.underline = True
# dotted underline | different types (https://python-docx.readthedocs.io/en/latest/api/enum/WdUnderline.html#wd-underline)
# j.font.underline = WD_UNDERLINE.DOT_DASH
# j.font.color.rgb = RGBColor(142, 131, 88) #() |(https://coolors.co/generate)
# k.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1 #sktblue
# k.font.color.theme_color = MSO_THEME_COLOR.ACCENT_2 #red
# k.font.color.theme_color = MSO_THEME_COLOR.ACCENT_3 #light parrot green  
# k.font.color.theme_color = MSO_THEME_COLOR.ACCENT_4 # purple
# k.font.color.theme_color = MSO_THEME_COLOR.ACCENT_5 # Aqua
# k.font.color.theme_color = MSO_THEME_COLOR.ACCENT_6 # orange
# k.font.color.theme_color = MSO_THEME_COLOR.BACKGROUND_1 #2
# k.font.color.theme_color = MSO_THEME_COLOR.DARK_1 #2

# https://www.resume.com/sample/civil-engineer/



#####################
# Header Section
#####################

# playing with sections
sections = d.sections[0] # first_page section
header = sections.header
para = header.add_paragraph('')
# para.text = "checking"
# para.font.color.rgb = RGBColor(142, 131, 88)
# para.styles = d.styles["Header"]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
l = para.add_run('Muhammed Roshan P S ', 'Strong')
l.font.size = Pt(24)
l.font.color.theme_color = MSO_THEME_COLOR.DARK_1
l.font.underline = True

sections = d.sections[0]  # Get the first section of the document
header = section.header
nxt = header.add_paragraph("")
nxt.alignment = WD_ALIGN_PARAGRAPH.CENTER
# k = nxt.add_run(' linkdin ')
# k = nxt.add_hyperlink(text='foobar', url='http://github.com')
# add_hyperlink(paragraph, url='http://google.com', text='Google', color=None/'FF8822', underline=True/False):
add_hyperlink(nxt, url='http://google.com', text='Github |',color='24292e',underline=True)
add_hyperlink(nxt, url='http://google.com', text=' Linkdin |',color='0A66C2',underline=True)
add_hyperlink(nxt, url='http://mailto:muhammedroshanps@gmail.com', text=' muhammedroshanps@gmail.com |',color='2b3137',underline=True)
add_hyperlink(nxt, url='http://tel:9809031477', text=' +919809031477',color='5B51D8',underline=True)
insertHR(nxt)
# add_hyperlink(nxt, text=' | Instagram ', url='http://google.com')
# add_hyperlink(nxt, text=' | Wiki ', url='http://google.com')
# k.font.size = Pt(10)
# k.font.color.theme_color = MSO_THEME_COLOR.DARK_1

d.save("wow.docx")