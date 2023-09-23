from  docx import Document
from docx.shared import Inches,Pt
d = Document()
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH 

j = d.add_paragraph('Muhammed Roshan P S')
j.add_run('owk')
j.add_run('oo')
j.style = 'List Bullet'
# j.paragraph_format.left_indent = Inches(0.03)


# Creating a table object
# table = d.add_table(rows=1,cols=1)

# Adding heading in the 1st row of the table
# row = table.rows[0].cells
# row[0].text = 'Id'
# row[1].text = 'Name'
# row[0].text = 'nice'
# p = row[0].paragraphs[0]
# # j = row[1].add_paragraph('wow')
# p.alignmennt = WD_ALIGN_PARAGRAPH.CENTER


table = d.add_table(rows=1,cols=2)
# row_cells = table.add_rows().cells
row_cells = table.rows[0].cells
row_cells[0].width = Inches(1)
row_cells[0].height = Pt(1)
row_cells[0].text = 'amount'
row_cells[1].text = 'wow'
row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
f = row_cells[0].paragraphs[0].runs[0].font
f.name = 'PT Sarif'
f.size = Pt(15)
# font = paragraph.runs[0].font


# Adding data from the list to the table
# for id, name in data:
# para = d.add_paragraph('')

# Adding a row and then adding data in it.
# row = table.add_row().cells
	# Converting id to string as table can only take string input
# l = row[0].add_paragraph('nice')
# j = row[1].add_paragraph('wow')
# j.alignmennt = WD_ALIGN_PARAGRAPH.RIGHT
# table.direction = WD_TABLE_DIRECTION.RTL
# row0._tr.addnext(row1._tr)

# Now save the document to a location
# doc.save('gfg.docx')



d.save('checktable.docx')




# from docx.enum.text import WD_ALIGN_PARAGRAPH

# def addCellText(row_cells, index, text):
#     row_cells[index].text = str(text)
#     paragraph=row_cells[index].paragraphs[0]
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
#     font = paragraph.runs[0].font
#     font.size= Pt(10)

# def addCellTextRight(row_cells, index, text):
#     row_cells[index].text = str(text)
#     paragraph=row_cells[index].paragraphs[0]
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#     font = paragraph.runs[0].font
#     font.size= Pt(10)