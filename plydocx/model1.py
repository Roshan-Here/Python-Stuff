from docx import Document
from docx.shared import Pt,Inches # for font sizing...
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor # for font color setting....
from docx.enum.dml import MSO_THEME_COLOR #MSO theme selector color
from docx.enum.text import WD_ALIGN_PARAGRAPH # alignmennt
from docx.enum.style import WD_STYLE
from functions import insertHR,add_hyperlink
from docx.enum.text import WD_LINE_SPACING # line spacing

#add a hyperlink with a custom color and no underline
# hyperlink = add_hyperlink(p, 'http://www.google.com', 'Google', 'FF8822', False)


###############################
# https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
################################

d = Document()

# for para creation simplicity 
def paara(spacing=False,Bbullet=False):
    para = d.add_paragraph('')
    if spacing == True:
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
    if Bbullet == 1:
        para.style = 'List Bullet'
    return para

# to setup font clr,theme,name for simpliity
# fff(dname=l,size=int(val),theme ='MSO_THEME_COLOR.ACCENT_1',fntname='Open Sans')
def fff(dname,size=int(11),theme=MSO_THEME_COLOR.ACCENT_1,fntname='Open Sans',bbold=False,iitalic=False):
    dname.font.size = Pt(size)
    dname.font.color.theme_color = theme
    if bbold == True:
        dname.bold = True
    else:
        pass
    if iitalic == 1:
        dname.italic = True
    else:
        pass
    dname.font.name = fntname

def addBlank(dname,count):
    for _ in range(count):
        dname.add_text(" ")


#####################
# Header Section
#####################


section = d.sections[0]  # Get the first section of the document
section.left_margin = Inches(.5)  # Set the left margin to 1 inch
section.right_margin = Inches(.5)  # Set the right margin to 1 inch
section.top_margin = Inches(.5)   # Set the top margin to 1 inch
section.bottom_margin = Inches(.5)

# playing with sections
sections = d.sections[0] # first_page section
header = sections.header
para = header.add_paragraph('')
# para.text = "checking"
# para.font.color.rgb = RGBColor(142, 131, 88)
# para.styles = d.styles["Header"]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
l = para.add_run('Muhammed Roshan P S', 'Strong')
l.font.size = Pt(24)
l.font.color.theme_color = MSO_THEME_COLOR.DARK_1
l.font.underline = True

sections = d.sections[0]  # Get the first section of the document
header = section.header
nxt = header.add_paragraph("")
nxt.alignment = WD_ALIGN_PARAGRAPH.CENTER
# k = nxt.add_run(' linkdin ')
# k = nxt.add_hyperlink(text='foobar', url='http://github.com')
# add_hyperlink(paragraph, url='http://google.com', text='Google', color=None/'FF8822', underline=True/False):
add_hyperlink(nxt, url='https://github.com/Roshan-Here', text='Github |',color='24292e',underline=True)
add_hyperlink(nxt, url='https://www.linkedin.com/in/muhammed-roshan-ps/', text=' Linkdin |',color='0A66C2',underline=True)
add_hyperlink(nxt, url='http://mailto:muhammedroshanps@gmail.com', text=' muhammedroshanps@gmail.com |',color='2b3137',underline=True)
add_hyperlink(nxt, url='http://tel:9809031477', text=' +919809031477',color='5B51D8',underline=True)
# insertHR(nxt,fontsize='20')
# add_hyperlink(nxt, text=' | Instagram ', url='http://google.com')
# add_hyperlink(nxt, text=' | Wiki ', url='http://google.com')
# k.font.size = Pt(10)
# k.font.color.theme_color = MSO_THEME_COLOR.DARK_1


##############################################
# Main Body
##############################################

# HEADING-1

para = d.add_heading('',level=5)
l = para.add_run('TECHNICAL SKILLS')
l.font.size = Pt(12)
l.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
l.font.name = 'Calibri'
insertHR(para,color='black')
# below the heading....
para = paara(spacing=True) # mov on to simplicity
l = para.add_run('   Languages   : ')
j = para.add_run('Python,C,JavaScript...etc')
# para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
fff(l,size=int(12),theme=MSO_THEME_COLOR.DARK_1,fntname='Calibri',bbold=True)
para = paara(spacing=True)
l = para.add_run('   Database     : ')
j = para.add_run('MySql,MongoDb...etc') # change... later
# para.paragraph_format.line_spacing = 0.5
fff(l,size=int(12),theme=MSO_THEME_COLOR.DARK_1,fntname='Calibri',bbold=True)
para = paara(spacing=True)
l = para.add_run('   Libraries      : ')
j = para.add_run('BeautifulSoup,Pyrogram,...etc') # change... later
fff(l,size=int(12),theme=MSO_THEME_COLOR.DARK_1,fntname='Calibri',bbold=True)
para = paara(spacing=True)
l = para.add_run('   Dev Tools    : ')
j = para.add_run('Visual Studio Code, Pycharm, Github, Vim, OpenAI') # change... later
fff(l,size=int(12),theme=MSO_THEME_COLOR.DARK_1,fntname='Calibri',bbold=True)


# HEADING-2

para = d.add_heading('',level=5)
l = para.add_run('EDUCATION')
l.font.size = Pt(12)
l.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
l.font.name = 'Calibri'
insertHR(para,color='black')
para = paara(spacing=True)
l = para.add_run('   SNMIMT COLLEGE OF ENGINEERING')
addBlank(l,68) # count depends on location string.len 
j = para.add_run('Kerala,India')
fff(l,size=int(12),bbold=True,theme=MSO_THEME_COLOR.DARK_1,fntname='Tahoma')
para = paara(spacing=True)
l = para.add_run('    Bachelor of Science in Computer Science')
fff(l,size=int(10),iitalic=True,fntname='Calibri',theme=MSO_THEME_COLOR.DARK_1)
para = paara()
l = para.add_run('    Ongoing (Expected Graduation: May 2024)') #ask if its compled/ongoing on form 
fff(l,size=int(10),iitalic=True,fntname='Calibri',theme=MSO_THEME_COLOR.DARK_1)


para = paara(spacing=True)
l = para.add_run('   Govt Boys Aluva')
addBlank(l,103)
j = para.add_run('Kerala,India')
fff(l,size=int(12),bbold=True,theme=MSO_THEME_COLOR.DARK_1,fntname='Tahoma')
para = paara(spacing=True)
l = para.add_run('    XII - Bio Maths - Science Grad.')
fff(l,size=int(10),iitalic=True,fntname='Calibri',theme=MSO_THEME_COLOR.DARK_1)


# HEADING -3

para = d.add_heading('',level=5)
l = para.add_run('PROJECTS')
l.font.size = Pt(12)
l.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
l.font.name = 'Calibri'
insertHR(para,color='black')

para = paara(spacing=True)
l = para.add_run('   Students Corner | ')
j = para.add_run('Collage Internel Manager')
fff(l,size=int(12),bbold=True,theme=MSO_THEME_COLOR.DARK_1,fntname='Tahoma')
fff(j,size=int(12),bbold=True,theme=MSO_THEME_COLOR.DARK_1,fntname='Cascadia MonoCascadia Mono')
para = paara(spacing=True)
add_hyperlink(para, url='https://github.com/Roshan-Here', text='    StudentsCorner.com',color='0A66C2',underline=True)

for _ in range(3):
    para = paara(spacing=True)
    l = para.add_run('       •')
    l.add_text(' An Web app to manage collage Students internel mark and attendance')
    fff(l,size=int(10),iitalic=True,fntname='Calibri',theme=MSO_THEME_COLOR.DARK_1)


para = paara(spacing=True)
l = para.add_run('   Students Corner | ')
j = para.add_run('Collage Internel Manager')
fff(l,size=int(12),bbold=True,theme=MSO_THEME_COLOR.DARK_1,fntname='Tahoma')
fff(j,size=int(12),bbold=True,theme=MSO_THEME_COLOR.DARK_1,fntname='Cascadia MonoCascadia Mono')
para = paara(spacing=True)
add_hyperlink(para, url='https://github.com/Roshan-Here', text='    StudentsCorner.com',color='0A66C2',underline=True)

for _ in range(3):
    para = paara(spacing=True)
    l = para.add_run('       •')
    l.add_text(' An Web app to manage collage Students internel mark and attendance')
    fff(l,size=int(10),iitalic=True,fntname='Calibri',theme=MSO_THEME_COLOR.DARK_1)
# i'm out of content...



# https://www.resume.com/sample/civil-engineer/

# work --tomorrow | sleep well 


d.save("wow.docx")