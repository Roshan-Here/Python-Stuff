from docx import Document
from docx.shared import Pt,Inches,Cm # for font sizing...
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor # for font color setting....
from docx.enum.dml import MSO_THEME_COLOR #MSO theme selector color
from docx.enum.text import WD_ALIGN_PARAGRAPH # alignmennt
from docx.enum.style import WD_STYLE
from functions import insertHR,add_hyperlink,fff
from docx.enum.text import WD_LINE_SPACING # line spacing
from docx.enum.table import WD_ROW_HEIGHT_RULE


d = Document()
def paara(spacing=False,Bbullet=False):
    para = d.add_paragraph('')
    if spacing == True:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(10)
    if Bbullet == 1:
        para.style = 'List Bullet'
    return para


#####################
# Header Section
#####################

# headding-1
section = d.sections[0]  # Get the first section of the document
section.left_margin = Inches(.5)  # Set the left margin to 1 inch
section.right_margin = Inches(.5)  # Set the right margin to 1 inch
section.top_margin = Inches(.5)   # Set the top margin to 1 inch
section.bottom_margin = Inches(.5)

sections = d.sections[0] # first_page section
header = sections.header
para = header.add_paragraph('')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
l = para.add_run('Muhammed Roshan P S')
l.bold = True
l.font.size = Pt(19)
l.font.name = 'PT Serif'
l.font.color.theme_color = MSO_THEME_COLOR.DARK_1
# l.font.underline = True


# sections = d.sections[0]  # Get the first section of the document
# header = section.header
nxt = header.add_paragraph("")
nxt.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_hyperlink(nxt, url='http://mailto:muhammedroshanps@gmail.com', text=' muhammedroshanps@gmail.com |',color='2b3137',underline=True)
add_hyperlink(nxt, url='https://github.com/Roshan-Here', text='Github |',color='24292e',underline=True)
add_hyperlink(nxt, url='https://www.linkedin.com/in/muhammed-roshan-ps/', text=' Linkdin |',color='0A66C2',underline=True)
add_hyperlink(nxt, url='http://tel:9809031477', text=' +919809031477',color='5B51D8',underline=True)
# nxt.font.name = 'PT Serif'
# nxt.font.size = Pt(11)
insertHR(nxt,fontsize='20',typee='thick')

###############
#main_body
###############

para = d.add_heading('',level=5)
l = para.add_run('Summary')
l.font.size = Pt(11)
l.font.name = 'PT Serif'
l.bold = True
l.font.color.theme_color = MSO_THEME_COLOR.DARK_1 

para = paara(spacing=True)
j = para.add_run('Financial Advisor with 7+ years of experience delivering financial/investment advisory services to high value clients. Proven success in managing multi-million dollar portfolios, driving profitability, and increasing ROI through skillful strategic planning, consulting, and financial advisory services.')
# para.paragraph_format.line_spacing = Pt(1)
fff(j,size=int(10),fntname='PT Serif')

#heading -2
para = d.add_heading('',level=5)
j = para.add_run('Professional Experience')
fff(j,size=int(11),fntname='PT Serif',bbold=True)


# para = paara(spacing=True)
# j = para.add_run('Month YYYY-Present')
# para.alignment =  WD_ALIGN_PARAGRAPH.RIGHT
# fff(j,size=int(10),fntname='PT Serif',iitalic=True)

table = d.add_table(rows=1,cols=2)
# row_cells = table.add_rows().cells
row_cells = table.rows[0].cells
row_cells[0].width = Inches(10)
row_cells[0].height = Cm(0.01)
# row_cells[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
row_cells[0].text = 'Job Heading'
row_cells[1].text = 'Month YYYY-Present'
row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
f = row_cells[0].paragraphs[0].runs[0].font
y = row_cells[1].paragraphs[0].runs[0].font
f.name = 'PT Sarif'
f.size = Pt(10)

y.name = 'PT Sarif'
y.size = Pt(10)
y.italic = True
para = paara(spacing=True)
k = para.add_run('Company name,place')
fff(k,size=int(10),fntname='PT Serif')






# para = paara(spacing=True)
# j = para.add_run('Job Heading')
# k = j.add_text('Month YYYY-Present')
# k.alignment = WD_ALIGN_PARAGRAPH.RIGHT
# fff(j,size=int(10),fntname='PT Serif',bbold=True)
# para = paara()
# k = para.add_run('Company name,place')
# fff(k,size=int(10),fntname='PT Serif')

for _ in range(3):
    para = paara()
    l = para.add_run('       •')
    l.add_text(' Learned applicable usage of Djago applications')
    fff(l,size=int(10),iitalic=True,fntname='Calibri',theme=MSO_THEME_COLOR.DARK_1)





d.save('nice.docx')